import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---- Page margins ----
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ---- Helper functions ----
def add_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, italic=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F5C38')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = str(val)
            for run in row.cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    return table

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def page_break():
    doc.add_page_break()

# ================================================================
# TITLE PAGE
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Rice Map Dashboard")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 92, 56)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("สรุปการสนทนาและการพัฒนาระบบ\nSession Log & Development Summary")
run2.font.size = Pt(14)
run2.italic = True
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run("20 เมษายน 2568  |  บันทึกโดย Antigravity AI")
date_run.font.size = Pt(11)
date_run.font.color.rgb = RGBColor(120, 120, 120)

url_p = doc.add_paragraph()
url_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
url_run = url_p.add_run("https://ton-munoi99.github.io/rice-map/")
url_run.font.size = Pt(11)
url_run.font.color.rgb = RGBColor(31, 92, 56)

page_break()

# ================================================================
# SECTION 1: OVERVIEW
# ================================================================
add_heading("1. ภาพรวมโปรเจกต์ (Project Overview)", 1, (31, 92, 56))
add_para(
    "Rice Map Dashboard คือเว็บแอปพลิเคชันแบบ Single Page Application (SPA) แสดงแผนที่ Choropleth "
    "ข้าวนาปีไทยรายจังหวัด ครอบคลุม 77 จังหวัด โดยไม่ใช้ Framework ภายนอกใดๆ — สร้างด้วย "
    "HTML + Vanilla CSS + Vanilla JavaScript ทั้งหมดอยู่ในไฟล์ index.html ไฟล์เดียว",
    size=11
)

doc.add_paragraph()
add_heading("ลิงก์สำคัญ", 2)
add_bullet("Live Website: https://ton-munoi99.github.io/rice-map/")
add_bullet("GitHub Repository: https://github.com/Ton-Munoi99/rice-map")
add_bullet("Workspace: C:\\Users\\sponlapatp\\Desktop\\Rice Map")

doc.add_paragraph()
add_heading("Tech Stack", 2)
add_table(
    ["ส่วน", "รายละเอียด"],
    [
        ["UI / Logic", "Vanilla HTML + CSS + JavaScript (ไม่มี Framework)"],
        ["แผนที่ (Map)", "SVG Choropleth จาก thailand-data.js (GeoJSON → SVG Paths)"],
        ["ข้อมูล", "rice-data.js → window.RICE_DATA_ROWS + window.PROVINCE_HOUSEHOLDS"],
        ["Fonts", "Cormorant Garamond + IBM Plex Mono + Sarabun (Google Fonts)"],
        ["Hosting", "GitHub Pages (ฟรี)"],
        ["Automation", "GitHub Actions (cron job อัตโนมัติ)"],
        ["ไม่ใช้", "React, Vue, D3.js, Leaflet, Tailwind CSS, shadcn"],
    ]
)

page_break()

# ================================================================
# SECTION 2: DATA LAYERS
# ================================================================
add_heading("2. ชั้นข้อมูล (Data Layers) ปัจจุบัน", 1, (31, 92, 56))
add_para("Dashboard มี 5 Layer ให้เลือกแสดงบนแผนที่:", size=11)
doc.add_paragraph()
add_table(
    ["Layer Key", "ชื่อภาษาไทย", "ข้อมูล", "Unit"],
    [
        ["production", "ผลผลิต", "OAE 2565-2567 + ประมาณการ 2568-2569", "ตัน/ปี"],
        ["yield", "ผลผลิตต่อไร่", "OAE 2565-2567 + ประมาณการ 2568-2569", "กก./ไร่"],
        ["area", "เนื้อที่เก็บเกี่ยว", "OAE 2565-2567 + ประมาณการ 2568-2569", "ไร่"],
        ["price", "ราคาที่โรงสีรับซื้อ", "สมาคมโรงสีข้าวไทย 2568-2569 (รายจังหวัด)", "บาท/ตัน"],
        ["households", "ครัวเรือนเกษตรกร", "OAE 2566 (ทุกสินค้าเกษตร)", "ครัวเรือน"],
    ]
)

doc.add_paragraph()
add_heading("สถานะข้อมูลรายปี", 2)
add_table(
    ["ปี", "Production/Yield/Area", "ราคา", "หมายเหตุ"],
    [
        ["2565", "✅ OAE จริง", "❌ ไม่มี", "OAE Yearbook"],
        ["2566", "✅ OAE จริง", "❌ ไม่มี", "OAE Yearbook"],
        ["2567", "✅ OAE จริง", "❌ ไม่มี", "OAE Yearbook"],
        ["2568", "📊 ประมาณการ Trend", "✅ สมาคมโรงสีข้าวไทย", "Trend + กรมการข้าว Calibration"],
        ["2569", "📊 ประมาณการ Trend", "✅ สมาคมโรงสีข้าวไทย", "Trend Estimate Only"],
    ]
)

page_break()

# ================================================================
# SECTION 3: DATA SOURCES
# ================================================================
add_heading("3. แหล่งข้อมูล (Data Sources)", 1, (31, 92, 56))
add_table(
    ["แหล่งข้อมูล", "ข้อมูล", "ปี", "URL/หมายเหตุ"],
    [
        ["OAE / สศก. (Yearbook)", "Production, Yield, Area รายจังหวัด", "2565–2567", "oae.go.th"],
        ["OAE Data Catalog", "ครัวเรือนเกษตรกรรายจังหวัด", "2566", "catalog.oae.go.th"],
        ["สมาคมโรงสีข้าวไทย (TRMA)", "ราคาข้าวเปลือกรายจังหวัด", "2568–2569", "thairicemillers.org"],
        ["TREA (สมาคมผู้ส่งออก)", "F.O.B. ราคาส่งออก", "รายสัปดาห์", "thairice.org"],
        ["กรมการข้าว (RD)", "เนื้อที่เพาะปลูก (Scale Factor)", "2568", "ricethailand.go.th"],
        ["Trend Estimation Script", "Production/Yield/Area ปี 2568–2569", "2568–2569", "scripts/estimate_2568_2569.js"],
    ]
)

doc.add_paragraph()
add_heading("Live Price Widget (3 ระดับราคา)", 2)
add_para("Widget ราคาแบบ Real-time มี 3 ชั้น:", bold=True)
add_bullet("🚢 F.O.B. (ราคาส่งออกหน้าท่าเรือ) — TREA อัพเดตรายสัปดาห์ → data/trea-fob.json")
add_bullet("🌾 OAE Farm-gate — ราคาหน้าไร่นาระดับประเทศ → data/prices-live.json")
add_bullet("🏭 ราคาที่โรงสีรับซื้อ — สมาคมโรงสีข้าวไทย รายจังหวัด → ฝังใน rice-data.js")

page_break()

# ================================================================
# SECTION 4: AUTOMATION
# ================================================================
add_heading("4. ระบบ Automation (GitHub Actions)", 1, (31, 92, 56))
add_para(
    "ใช้ GitHub Actions เป็น Automation Engine ดึงข้อมูลจากแหล่งภายนอกอัตโนมัติ "
    "แล้ว Commit กลับเข้า Repository โดยไม่ต้องทำมือ:", size=11
)
doc.add_paragraph()

add_heading("Workflow: Update TREA FOB Prices", 2)
add_bullet("ไฟล์: .github/workflows/update-trea-fob.yml")
add_bullet("เวลา: ทุกวัน พุธ และ พฤหัส (อัตโนมัติ)")
add_bullet("Script: scripts/fetch_trea_fob.py")
add_bullet("Output: data/trea-fob.json → Commit อัตโนมัติ → GitHub Pages Deploy")

doc.add_paragraph()
add_heading("ทำไมถึงใช้ GitHub ไม่ใช้ Netlify?", 2)
add_table(
    ["ฟีเจอร์", "GitHub Pages + Actions", "Netlify"],
    [
        ["Static Hosting", "✅ ฟรี", "✅ ฟรี (จำกัด)"],
        ["Scheduled Cron Jobs", "✅ ฟรี (GitHub Actions)", "❌ ต้องจ่ายเงิน"],
        ["รัน Python Script", "✅ ได้เลย", "❌ Build Only"],
        ["Commit ข้อมูลกลับ Repo", "✅ ทำได้", "❌ ไม่ได้"],
        ["CORS-free (ข้อมูลอบไว้)", "✅ ข้อมูลอยู่ใน repo", "❌ ต้อง call API จาก browser"],
        ["Version History ข้อมูล", "✅ Git log ทุก commit", "❌ ไม่มี"],
        ["ราคา (Public Repo)", "✅ ฟรีหมด", "⚠️ มีจำกัด"],
    ]
)

doc.add_paragraph()
add_heading("Pattern การดึง API ภายนอก (Best Practice)", 2)
add_para("ขั้นตอนการทำงานของระบบ:", bold=True)
add_bullet("1. GitHub Actions รัน Python Script ตาม Schedule (cron)")
add_bullet("2. Script ดึงข้อมูลจาก API ภายนอก (OAE, TREA, BoT, กรมอุตุฯ ฯลฯ)")
add_bullet("3. แปลงข้อมูลเป็น JSON/JS")
add_bullet("4. Commit + Push กลับ Repository อัตโนมัติ")
add_bullet("5. GitHub Pages Deploy อัตโนมัติ")
add_bullet("6. ผู้ใช้เห็นข้อมูลใหม่ทันที (ไม่ต้อง refresh เป็นพิเศษ)")

doc.add_paragraph()
add_para("ข้อดีของ Pattern นี้:", bold=True)
add_bullet("ข้อมูลถูก 'อบไว้ใน repo' → เว็บโหลดเร็ว ไม่ต้อง call API ทุกครั้ง")
add_bullet("ไม่มีปัญหา CORS หรือ API Rate Limit จาก Browser")
add_bullet("ข้อมูลมี Version History — ย้อนดูว่าราคาวันไหนเท่าไหร่ได้")
add_bullet("API ภายนอก down ชั่วคราว → ข้อมูลเก่ายังแสดงได้ ไม่พัง")

page_break()

# ================================================================
# SECTION 5: FILE STRUCTURE
# ================================================================
add_heading("5. โครงสร้างไฟล์ (File Structure)", 1, (31, 92, 56))
add_table(
    ["ไฟล์/โฟลเดอร์", "หน้าที่"],
    [
        ["index.html", "Main App (~2,668 บรรทัด) — HTML + CSS + JS ทั้งหมด"],
        ["rice-data.js", "window.RICE_DATA_ROWS + window.PROVINCE_HOUSEHOLDS"],
        ["thailand-data.js", "GeoJSON SVG Paths 77 จังหวัด (~1MB)"],
        ["rice-data.csv", "CSV Snapshot สำรอง"],
        ["SESSION_LOG.md", "บันทึกการสนทนาทั้งหมด"],
        ["data/trea-fob.json", "ราคา F.O.B. (TREA) — อัพเดตอัตโนมัติ"],
        ["data/prices-live.json", "OAE Farm-gate Prices"],
        ["data/farmer_households.csv", "จำนวนครัวเรือนเกษตรกรรายจังหวัด OAE 2566"],
        ["scripts/fetch_trea_fob.py", "Scraper ราคา FOB จาก TREA"],
        ["scripts/fetch_oae_prices.py", "ดึงราคา OAE Farm-gate"],
        ["scripts/build_households.py", "สร้าง window.PROVINCE_HOUSEHOLDS จาก CSV"],
        ["scripts/estimate_2568_2569.js", "คำนวณ Trend Estimate 2568–2569"],
        ["scripts/check_syntax.js", "ตรวจ JS Syntax ของ index.html"],
        [".github/workflows/update-trea-fob.yml", "GitHub Action อัพเดต FOB ทุก พุธ-พฤหัส"],
    ]
)

page_break()

# ================================================================
# SECTION 6: HOUSEHOLD DATA
# ================================================================
add_heading("6. ข้อมูลครัวเรือนเกษตรกร (OAE 2566)", 1, (31, 92, 56))
add_para(
    "เพิ่ม Layer 'ครัวเรือนเกษตรกร' ใหม่ โดยดาวน์โหลดข้อมูลจาก OAE Data Catalog "
    "ปี 2566 ครอบคลุมทุกสินค้าเกษตร (ไม่ใช่เฉพาะข้าว)", size=11
)
doc.add_paragraph()
add_table(
    ["แหล่งข้อมูล", "รายละเอียด"],
    [
        ["URL", "catalog.oae.go.th/dataset/e21cc6cd-8641-44c9-b335-eceb548b83b4"],
        ["ไฟล์", "data/farmer_households.csv (Encoding: CP874)"],
        ["Script", "scripts/build_households.py"],
        ["Output", "window.PROVINCE_HOUSEHOLDS ใน rice-data.js"],
        ["ปีข้อมูล", "2566 (ล่าสุดที่ OAE เผยแพร่)"],
        ["Coverage", "76 จังหวัด"],
    ]
)

doc.add_paragraph()
add_heading("ตัวอย่างข้อมูล (Top 5 จังหวัด)", 2)
add_table(
    ["จังหวัด", "ครัวเรือนเกษตรกร"],
    [
        ["นครราชสีมา", "366,974"],
        ["อุบลราชธานี", "348,062"],
        ["ขอนแก่น", "265,770"],
        ["ศรีสะเกษ", "267,646"],
        ["สุรินทร์", "251,970"],
    ]
)

doc.add_paragraph()
add_heading("Bug ที่พบและแก้ไขแล้ว", 2)
add_para("ปัญหา:", bold=True)
add_para(
    "  const PROVINCE_HOUSEHOLDS = {...} ใน rice-data.js ไม่ผูกกับ window "
    "ทำให้ index.html เช็ค window.PROVINCE_HOUSEHOLDS แล้วได้ undefined",
    italic=True, color=(180, 0, 0)
)
add_para("วิธีแก้:", bold=True)
add_para(
    "  เปลี่ยนเป็น window.PROVINCE_HOUSEHOLDS = {...} "
    "เพื่อให้ Script อื่น Access ได้ข้ามไฟล์",
    italic=True, color=(0, 120, 0)
)

page_break()

# ================================================================
# SECTION 7: NEXT STEPS
# ================================================================
add_heading("7. แผนงานต่อไป (Next Steps)", 1, (31, 92, 56))
add_bullet("เพิ่มข้อมูลต้นทุนปุ๋ย / ต้นทุนการผลิตรายจังหวัด (OAE PDF — ยาก ต้อง OCR)")
add_bullet("เมื่อ OAE ออก Yearbook ปี 2568 → นำเข้าผ่าน Import CSV แทน estimated_trend")
add_bullet("พิจารณา Layer จำนวนเกษตรกรปลูกข้าวโดยเฉพาะ (DOAE Digital Farmer)")
add_bullet("เพิ่ม API อัตราแลกเปลี่ยน BoT (USD/THB) → แสดงราคา FOB เป็นบาทพร้อมกัน")
add_bullet("เพิ่มข้อมูลสภาพอากาศรายจังหวัด (กรมอุตุนิยมวิทยา API)")
add_bullet("Mobile Responsive ปรับปรุง Side Panel Collapse")

doc.add_paragraph()
add_heading("API ภายนอกที่น่าสนใจเพิ่มเติม", 2)
add_table(
    ["แหล่ง", "API", "ข้อมูล", "การเชื่อมต่อ"],
    [
        ["กรณีผ้าพาณิชย์ (MOC)", "dataapi.moc.go.th", "ราคาสินค้าเกษตร", "✅ มี API จริง"],
        ["ธนาคารแห่งประเทศไทย (BoT)", "api.bot.or.th", "อัตราแลกเปลี่ยน USD/THB", "✅ ง่ายมาก"],
        ["กรมอุตุนิยมวิทยา", "data.tmd.go.th", "ฝน / อากาศรายจังหวัด", "✅ มี API"],
        ["กรมการข้าว", "ricethailand.go.th", "พื้นที่เพาะปลูก", "⚠️ ต้อง Scrape"],
        ["DOAE (กรมส่งเสริม)", "data.doae.go.th", "จำนวนชาวนาขึ้นทะเบียน", "✅ มี API"],
    ]
)

doc.add_paragraph()
add_heading("คำสั่ง Useful", 2)
add_bullet("Rebuild household data: python scripts/build_households.py")
add_bullet("Recalculate trend: node scripts/estimate_2568_2569.js")
add_bullet("Check JS syntax: node scripts/check_syntax.js")
add_bullet("Deploy: git add -A && git commit -m '...' && git push")

page_break()

# ================================================================
# SECTION 8: COMMITS
# ================================================================
add_heading("8. GitHub Commits ล่าสุด (20 เม.ย. 2568)", 1, (31, 92, 56))
add_table(
    ["Commit", "รายละเอียด"],
    [
        ["58f0342", "docs: add SESSION_LOG.md - full conversation record"],
        ["6a0793c", "fix: assign PROVINCE_HOUSEHOLDS to window for cross-script access"],
        ["d5ab192", "feat: add farmer household layer (OAE 2566)"],
        ["bd23c49", "style: user reverted FOB price styling back to default size"],
        ["b372938", "style: emphasize FOB export prices"],
    ]
)

# ================================================================
# FOOTER
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    f"บันทึกโดย Antigravity AI  |  20 เมษายน 2568  |  Rice Map Project"
)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(150, 150, 150)
footer_run.italic = True

# ================================================================
# SAVE
# ================================================================
out_path = r"C:\Users\sponlapatp\Desktop\RiceMap_Summary_20Apr2568.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
